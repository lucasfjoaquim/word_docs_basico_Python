from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def criar_arquivo_word(nome_arquivo):
    doc = Document()
    paragrafo = doc.add_paragraph("""
    Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nunc lacus nunc, tincidunt quis tellus vel, tincidunt ornare magna. Phasellus in dignissim leo, quis fringilla lacus. Fusce ultricies tincidunt nisi ut laoreet. Quisque condimentum odio massa, a placerat mauris maximus eget. Fusce id leo quis arcu fringilla ornare. Sed aliquam malesuada semper. Suspendisse et orci vitae ante consequat vestibulum sed vitae ipsum.

    Fusce in volutpat mauris. Proin eget congue erat. Sed a laoreet tortor. Ut lobortis velit vel metus hendrerit, eu semper purus mollis. Fusce ut risus sit amet nunc volutpat pulvinar eu id est. Maecenas lobortis massa purus, eu consectetur neque hendrerit pulvinar. Nunc sit amet tincidunt neque, in egestas magna. Suspendisse potenti. Cras sed ante placerat, vestibulum arcu bibendum, aliquet ligula. Cras et dolor eu eros varius sagittis. Nullam eu ultrices elit.

    Donec diam est, auctor vel convallis sagittis, vestibulum at neque. Mauris interdum, dolor id vulputate porta, ipsum nibh congue nisl, ut vehicula enim dui et velit. Nulla vitae leo efficitur, imperdiet ante vel, dictum magna. Sed molestie sapien sem, vitae vulputate libero aliquet sed. Vestibulum eleifend pulvinar eros quis blandit. Curabitur volutpat turpis elementum est eleifend tincidunt. Praesent non tellus ornare, aliquet odio sed, consequat lorem. Nunc vel pellentesque mi. Nulla nec velit ante. Mauris fringilla, elit ac placerat elementum, nulla est accumsan purus, sed sollicitudin velit nisi ac justo. Sed nisl neque, interdum id rutrum malesuada, consectetur non tortor. Phasellus at dolor sed dolor tristique dictum non a magna. Donec vitae enim dui.

    Duis sed eros tincidunt, iaculis urna nec, tincidunt orci. Class aptent taciti sociosqu ad litora torquent per conubia nostra, per inceptos himenaeos. Maecenas et urna felis. Ut fermentum tellus sapien, sit amet lobortis ipsum vulputate et. Pellentesque ac nisi massa. Donec aliquet enim eu aliquam ornare. Vestibulum maximus et ligula interdum interdum. Curabitur a nunc lectus. Vivamus viverra dictum metus in congue. Suspendisse tristique lectus in ex maximus, eu iaculis leo auctor. Quisque congue eleifend bibendum. Aenean dapibus arcu vel nisl porttitor, nec finibus nunc fringilla. Duis non elit consectetur, pretium leo eget, commodo sem. Nulla fermentum sem sem, posuere venenatis libero dictum et.

    Donec tempus congue sodales. Sed feugiat faucibus risus at convallis. Morbi sem mauris, pellentesque vitae lacus nec, hendrerit rutrum eros. Donec dapibus libero mi, sit amet hendrerit metus rhoncus et. Ut vel posuere urna. Curabitur vel porta mauris. Nullam consequat malesuada tellus eu vehicula. Cras non nisi lacinia, lacinia ante a, porttitor eros. Praesent ullamcorper ipsum sed lobortis ornare. Nulla condimentum ipsum a suscipit sagittis. Cras varius ultricies elit et vestibulum. Cras malesuada mi at ex lobortis malesuada. Aenean sed augue mi. Donec feugiat, ante at tincidunt aliquam, ante risus placerat velit, et finibus lacus nisi quis ligula.""")
    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.save(nome_arquivo)

nome_arquivo = "arquivo_word.docx"

try:
    criar_arquivo_word(nome_arquivo)
    print(f"Arquivo '{nome_arquivo}' criado com sucesso.")
except FileExistsError:
    doc = Document(nome_arquivo)
    paragrafo = doc.add_paragraph("Lorem Ipsum")
    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.save(nome_arquivo)
    print(f"Texto 'Lorem Ipsum' inserido no arquivo '{nome_arquivo}' existente.")